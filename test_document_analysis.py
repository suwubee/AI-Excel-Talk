#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
文档分析功能测试脚本
用于验证新增的文档分析功能是否正常工作
"""

import os
import sys
from pathlib import Path

def test_imports():
    """测试模块导入"""
    print("🔧 测试模块导入...")
    
    try:
        from document_analyzer import DocumentAnalyzer
        print("✅ DocumentAnalyzer 导入成功")
    except ImportError as e:
        print(f"❌ DocumentAnalyzer 导入失败: {e}")
        return False
    
    try:
        from document_utils import AdvancedDocumentProcessor, DocumentSearchEngine
        print("✅ DocumentUtils 导入成功")
    except ImportError as e:
        print(f"❌ DocumentUtils 导入失败: {e}")
        return False
    
    try:
        from document_ai_analyzer import EnhancedDocumentAIAnalyzer
        print("✅ DocumentAIAnalyzer 导入成功")
    except ImportError as e:
        print(f"❌ DocumentAIAnalyzer 导入失败: {e}")
        return False
    
    return True

def test_dependencies():
    """测试依赖包"""
    print("\n📦 测试依赖包...")
    
    dependencies = [
        ('markitdown', 'MarkItDown'),
        ('docx', 'python-docx'),
        ('PyPDF2', 'PyPDF2'),
        ('fitz', 'PyMuPDF'),
        ('PIL', 'Pillow')
    ]
    
    missing_deps = []
    
    for module_name, package_name in dependencies:
        try:
            __import__(module_name)
            print(f"✅ {package_name} 可用")
        except ImportError:
            print(f"❌ {package_name} 缺失")
            missing_deps.append(package_name)
    
    if missing_deps:
        print(f"\n⚠️ 缺失的依赖包: {', '.join(missing_deps)}")
        print("请运行: pip install -r requirements.txt")
        return False
    
    return True

def test_document_analyzer():
    """测试文档分析器基本功能"""
    print("\n📊 测试文档分析器...")
    
    try:
        from document_analyzer import DocumentAnalyzer
        analyzer = DocumentAnalyzer()
        print("✅ DocumentAnalyzer 初始化成功")
        
        # 测试支持的文档类型
        supported_types = analyzer._detect_document_type("test.docx")
        if supported_types == "docx":
            print("✅ DOCX类型检测正常")
        
        supported_types = analyzer._detect_document_type("test.pdf")
        if supported_types == "pdf":
            print("✅ PDF类型检测正常")
        
        return True
        
    except Exception as e:
        print(f"❌ DocumentAnalyzer 测试失败: {e}")
        return False

def test_document_processor():
    """测试文档处理器"""
    print("\n🔧 测试文档处理器...")
    
    try:
        from document_utils import AdvancedDocumentProcessor
        processor = AdvancedDocumentProcessor()
        print("✅ AdvancedDocumentProcessor 初始化成功")
        
        # 测试支持的格式
        expected_formats = ['.docx', '.pdf']
        if processor.supported_formats == expected_formats:
            print("✅ 支持的文件格式配置正确")
        
        return True
        
    except Exception as e:
        print(f"❌ AdvancedDocumentProcessor 测试失败: {e}")
        return False

def test_ai_analyzer():
    """测试AI分析器（不需要真实API Key）"""
    print("\n🤖 测试AI分析器...")
    
    try:
        from document_ai_analyzer import EnhancedDocumentAIAnalyzer
        
        # 使用假的API Key进行初始化测试
        ai_analyzer = EnhancedDocumentAIAnalyzer("test_key", "https://api.openai.com/v1", "gpt-4o-mini")
        print("✅ EnhancedDocumentAIAnalyzer 初始化成功")
        
        # 测试任务建议功能
        fake_analysis = {
            'file_info': {'type': 'pdf', 'size_mb': 5},
            'structure_analysis': {'headings': {1: [{'text': 'Test'}]}}
        }
        
        suggestions = ai_analyzer.suggest_analysis_tasks(fake_analysis)
        if len(suggestions) > 0:
            print(f"✅ 分析任务建议生成成功 ({len(suggestions)}个建议)")
        
        return True
        
    except Exception as e:
        print(f"❌ EnhancedDocumentAIAnalyzer 测试失败: {e}")
        return False

def test_markitdown_functionality():
    """测试MarkItDown功能"""
    print("\n📝 测试MarkItDown功能...")
    
    try:
        from markitdown import MarkItDown
        md = MarkItDown()
        print("✅ MarkItDown 初始化成功")
        
        # 创建一个测试文本文件
        test_file = "test_document.txt"
        with open(test_file, 'w', encoding='utf-8') as f:
            f.write("这是一个测试文档\n\n用于验证MarkItDown功能")
        
        # 测试转换
        result = md.convert(test_file)
        if result and hasattr(result, 'text_content'):
            print("✅ MarkItDown 文档转换成功")
        
        # 清理测试文件
        os.remove(test_file)
        
        return True
        
    except Exception as e:
        print(f"❌ MarkItDown 测试失败: {e}")
        # 清理测试文件（如果存在）
        if os.path.exists("test_document.txt"):
            os.remove("test_document.txt")
        return False

def create_sample_documents():
    """创建示例文档用于测试"""
    print("\n📄 创建示例文档...")
    
    try:
        # 创建示例DOCX文档
        from docx import Document
        doc = Document()
        doc.add_heading('测试文档标题', 0)
        doc.add_heading('第一章 简介', level=1)
        doc.add_paragraph('这是一个示例段落，用于测试文档分析功能。')
        doc.add_heading('第二章 功能说明', level=1)
        doc.add_paragraph('本文档包含多个章节和段落，用于验证标题识别和结构分析。')
        doc.add_heading('2.1 子章节', level=2)
        doc.add_paragraph('这是一个子章节的内容。包含关键词：合同编号 ABC123。')
        
        doc.save('sample_document.docx')
        print("✅ 示例DOCX文档创建成功: sample_document.docx")
        
        return True
        
    except Exception as e:
        print(f"❌ 示例文档创建失败: {e}")
        return False

def test_full_workflow():
    """测试完整的文档分析工作流"""
    print("\n🔄 测试完整工作流...")
    
    if not os.path.exists('sample_document.docx'):
        print("❌ 示例文档不存在，请先运行create_sample_documents()")
        return False
    
    try:
        from document_analyzer import DocumentAnalyzer
        from document_utils import AdvancedDocumentProcessor
        
        # 1. 初始化分析器
        analyzer = DocumentAnalyzer()
        processor = AdvancedDocumentProcessor()
        
        # 2. 加载文档
        analysis_result = analyzer.analyze_document('sample_document.docx')
        print("✅ 文档加载和分析完成")
        
        # 3. 检查分析结果结构
        expected_keys = ['file_info', 'preview_data', 'structure_analysis', 'ai_analysis_data', 'search_capabilities']
        for key in expected_keys:
            if key in analysis_result:
                print(f"✅ {key} 数据结构正确")
            else:
                print(f"❌ {key} 数据缺失")
        
        # 4. 测试搜索功能
        search_results = analyzer.search_keyword_context("合同编号", context_lines=2)
        if search_results:
            print(f"✅ 关键词搜索功能正常 (找到{len(search_results)}个结果)")
        else:
            print("⚠️ 关键词搜索无结果（这可能是正常的）")
        
        # 5. 测试代码生成
        code = analyzer.generate_analysis_code("搜索文档中的关键信息")
        if code and len(code) > 100:
            print("✅ 代码生成功能正常")
        
        return True
        
    except Exception as e:
        print(f"❌ 完整工作流测试失败: {e}")
        return False

def cleanup():
    """清理测试文件"""
    print("\n🧹 清理测试文件...")
    
    test_files = ['sample_document.docx', 'test_document.txt']
    for file in test_files:
        if os.path.exists(file):
            os.remove(file)
            print(f"✅ 已删除: {file}")

def main():
    """主测试函数"""
    print("🚀 文档分析功能测试开始")
    print("=" * 50)
    
    # 测试步骤
    tests = [
        ("模块导入", test_imports),
        ("依赖包检查", test_dependencies),
        ("文档分析器", test_document_analyzer),
        ("文档处理器", test_document_processor),
        ("AI分析器", test_ai_analyzer),
        ("MarkItDown功能", test_markitdown_functionality),
    ]
    
    results = []
    
    for test_name, test_func in tests:
        print(f"\n{'='*20} {test_name} {'='*20}")
        try:
            result = test_func()
            results.append((test_name, result))
        except Exception as e:
            print(f"❌ {test_name} 测试出现异常: {e}")
            results.append((test_name, False))
    
    # 如果基础测试都通过，进行高级测试
    if all(result for _, result in results):
        print(f"\n{'='*20} 高级测试 {'='*20}")
        
        # 创建示例文档
        if create_sample_documents():
            # 完整工作流测试
            workflow_result = test_full_workflow()
            results.append(("完整工作流", workflow_result))
        
        # 清理测试文件
        cleanup()
    
    # 总结结果
    print("\n" + "="*50)
    print("📊 测试结果总结")
    print("="*50)
    
    passed = 0
    total = len(results)
    
    for test_name, result in results:
        status = "✅ 通过" if result else "❌ 失败"
        print(f"{test_name:<20} {status}")
        if result:
            passed += 1
    
    print(f"\n总体结果: {passed}/{total} 测试通过")
    
    if passed == total:
        print("🎉 所有测试通过！文档分析功能已准备就绪。")
    else:
        print("⚠️ 部分测试失败，请检查依赖安装和配置。")
    
    return passed == total

if __name__ == "__main__":
    success = main()
    sys.exit(0 if success else 1) 