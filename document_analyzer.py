#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
文档智能分析器 - 支持docx、pdf分析
- 使用markitdown进行数据预览和清洗
- 结构化分析文档层级、段落、字体信息
- AI提示词友好输出
- 支持关键词搜索和上下文提取
"""

import os
import re
from typing import Dict, List, Any, Tuple, Optional
from pathlib import Path
import json

# 文档处理相关imports
try:
    from markitdown import MarkItDown
    MARKITDOWN_AVAILABLE = True
except ImportError as e:
    MarkItDown = None
    MARKITDOWN_AVAILABLE = False
    print(f"⚠️ markitdown 未安装或缺少依赖: {e}")
    print("📝 解决方案: pip install markitdown[all]")

try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    DOCX_AVAILABLE = True
except ImportError as e:
    Document = None
    DOCX_AVAILABLE = False
    print(f"⚠️ python-docx 未安装: {e}")
    print("📝 解决方案: pip install python-docx")

try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError as e:
    fitz = None
    PYMUPDF_AVAILABLE = False
    print(f"⚠️ PyMuPDF 未安装: {e}")
    print("📝 解决方案: pip install pymupdf")

try:
    import PyPDF2
    PYPDF2_AVAILABLE = True
except ImportError as e:
    PyPDF2 = None
    PYPDF2_AVAILABLE = False
    print(f"⚠️ PyPDF2 未安装: {e}")
    print("📝 解决方案: pip install PyPDF2")

# 可选依赖
try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    pdfplumber = None
    PDFPLUMBER_AVAILABLE = False

try:
    import docx2txt
    DOCX2TXT_AVAILABLE = True
except ImportError:
    docx2txt = None
    DOCX2TXT_AVAILABLE = False

from PIL import Image
import io
import base64

class DocumentAnalyzer:
    """文档智能分析器 - 支持docx、pdf文档的深度分析"""
    
    def __init__(self):
        self.document_path = None
        self.document_type = None
        self.analysis_result = {}
        self.markdown_converter = MarkItDown() if MarkItDown else None
        
        # 检查依赖可用性
        self._check_dependencies()
    
    def _check_dependencies(self) -> Dict[str, bool]:
        """检查所有依赖的可用性"""
        dependencies = {
            'markitdown': MARKITDOWN_AVAILABLE,
            'python-docx': DOCX_AVAILABLE,
            'pymupdf': PYMUPDF_AVAILABLE,
            'pypdf2': PYPDF2_AVAILABLE,
            'pdfplumber': PDFPLUMBER_AVAILABLE,
            'docx2txt': DOCX2TXT_AVAILABLE
        }
        return dependencies
    
    def get_missing_dependencies(self) -> List[str]:
        """获取缺失的关键依赖"""
        deps = self._check_dependencies()
        missing = []
        
        if not deps['markitdown']:
            missing.append('markitdown[all]')
        if not deps['python-docx']:
            missing.append('python-docx')
        if not deps['pymupdf'] and not deps['pypdf2']:
            missing.append('pymupdf 或 PyPDF2')
        
        return missing
    
    def is_ready_for_analysis(self, file_type: str = None) -> Tuple[bool, List[str]]:
        """检查是否准备好进行分析"""
        missing = self.get_missing_dependencies()
        
        if file_type == 'docx' and not DOCX_AVAILABLE:
            return False, ['python-docx']
        elif file_type == 'pdf' and not (PYMUPDF_AVAILABLE or PYPDF2_AVAILABLE):
            return False, ['pymupdf 或 PyPDF2']
        elif not MARKITDOWN_AVAILABLE:
            return False, ['markitdown[all]']
        
        return len(missing) == 0, missing
    
    def analyze_document(self, file_path: str) -> Dict[str, Any]:
        """
        全面分析文档
        
        Args:
            file_path: 文档文件路径
            
        Returns:
            完整的分析结果字典
        """
        self.document_path = file_path
        self.document_type = self._detect_document_type(file_path)
        
        if self.document_type not in ['docx', 'pdf']:
            raise ValueError(f"不支持的文件类型: {self.document_type}")
        
        # 检查依赖是否满足
        is_ready, missing_deps = self.is_ready_for_analysis(self.document_type)
        if not is_ready:
            raise ImportError(f"缺少必要依赖: {', '.join(missing_deps)}. 请运行: pip install {' '.join(missing_deps)}")
        
        analysis = {
            'file_info': {
                'path': file_path,
                'name': Path(file_path).name,
                'type': self.document_type,
                'size_mb': round(os.path.getsize(file_path) / (1024 * 1024), 2)
            },
            'preview_data': {},
            'structure_analysis': {},
            'ai_analysis_data': {},
            'search_capabilities': {}
        }
        
        # 1. 数据预览阶段 - 使用markitdown
        analysis['preview_data'] = self._generate_preview_with_markitdown(file_path)
        
        # 2. 结构化分析 - 从原始文档提取
        analysis['structure_analysis'] = self._analyze_document_structure(file_path)
        
        # 3. AI分析数据 - 为AI提供结构化信息
        analysis['ai_analysis_data'] = self._prepare_ai_analysis_data(analysis)
        
        # 4. 搜索功能准备
        analysis['search_capabilities'] = self._prepare_search_capabilities(file_path)
        
        self.analysis_result = analysis
        return analysis
    
    def _detect_document_type(self, file_path: str) -> str:
        """检测文档类型"""
        ext = Path(file_path).suffix.lower()
        if ext == '.docx':
            return 'docx'
        elif ext == '.pdf':
            return 'pdf'
        else:
            return 'unknown'
    
    def _generate_preview_with_markitdown(self, file_path: str) -> Dict[str, Any]:
        """
        使用markitdown生成预览数据（清除格式，保留图片，限制10页）
        """
        preview_data = {
            'status': 'success',
            'content': '',
            'images': [],
            'page_count': 0,
            'word_count': 0,
            'error': None,
            'has_images': False,
            'images_preview': []
        }
        
        try:
            if not self.markdown_converter:
                preview_data['status'] = 'error'
                preview_data['error'] = 'MarkItDown 未可用'
                return preview_data
            
            # 使用markitdown转换
            result = self.markdown_converter.convert(file_path)
            content = result.text_content if hasattr(result, 'text_content') else str(result)
            
            # 限制内容长度（模拟10页）
            max_chars = 20000  # 约10页内容
            if len(content) > max_chars:
                content = content[:max_chars] + "\n\n... [内容被截断，超过10页预览限制]"
            
            preview_data['content'] = content
            preview_data['word_count'] = len(content.split())
            
            # 估算页数（基于字符数）
            chars_per_page = 2000
            preview_data['page_count'] = max(1, len(content) // chars_per_page)
            
            # 提取图片信息用于预览
            images_info = self._extract_images_info(file_path)
            preview_data['images'] = [img['base64'] for img in images_info[:5]]  # 前5张图片的base64
            preview_data['images_preview'] = images_info[:5]  # 图片详细信息
            preview_data['has_images'] = len(images_info) > 0
            
        except Exception as e:
            preview_data['status'] = 'error'
            preview_data['error'] = f'MarkItDown处理失败: {str(e)}'
            
            # 尝试备用方案
            try:
                fallback_content = self._fallback_content_extraction(file_path)
                if fallback_content:
                    preview_data['content'] = fallback_content[:20000]
                    preview_data['word_count'] = len(fallback_content.split())
                    preview_data['status'] = 'fallback'
                    preview_data['error'] = f'使用备用方案: {str(e)}'
            except:
                preview_data['content'] = '内容提取失败，请检查文件格式和依赖安装'
        
        return preview_data
    
    def _fallback_content_extraction(self, file_path: str) -> str:
        """备用内容提取方案"""
        content = ""
        
        try:
            if self.document_type == 'docx' and Document:
                doc = Document(file_path)
                content = "\n".join([para.text for para in doc.paragraphs])
            elif self.document_type == 'pdf' and fitz:
                doc = fitz.open(file_path)
                for page_num in range(min(10, len(doc))):  # 限制10页
                    page = doc[page_num]
                    content += page.get_text() + "\n"
                doc.close()
        except:
            pass
        
        return content
    
    def _analyze_document_structure(self, file_path: str) -> Dict[str, Any]:
        """
        从原始文档提取结构化信息
        不使用markitdown的结果，直接从文档格式中提取
        """
        if self.document_type == 'docx':
            return self._analyze_docx_structure(file_path)
        elif self.document_type == 'pdf':
            return self._analyze_pdf_structure(file_path)
        else:
            return {'error': f'不支持的文档类型: {self.document_type}'}
    
    def _analyze_docx_structure(self, file_path: str) -> Dict[str, Any]:
        """分析DOCX文档结构"""
        try:
            if Document is None:
                return {'error': 'python-docx未安装'}
            
            doc = Document(file_path)
            
            structure = {
                'total_paragraphs': len(doc.paragraphs),
                'headings': {},
                'fonts_used': set(),
                'paragraph_analysis': [],
                'tables_count': len(doc.tables),
                'images_count': 0,
                'page_breaks': 0,
                'images_info': [],
                'watermark_analysis': {}
            }
            
            # 分析段落和标题层级
            heading_levels = {}
            for i, para in enumerate(doc.paragraphs):
                para_info = {
                    'index': i,
                    'text': para.text.strip()[:100],  # 前100字符
                    'style': para.style.name if para.style else 'Normal',
                    'alignment': str(para.alignment) if para.alignment else 'LEFT',
                    'is_heading': False,
                    'heading_level': 0
                }
                
                # 检测标题
                if para.style and 'Heading' in para.style.name:
                    para_info['is_heading'] = True
                    level_match = re.search(r'Heading (\d+)', para.style.name)
                    if level_match:
                        level = int(level_match.group(1))
                        para_info['heading_level'] = level
                        
                        if level not in heading_levels:
                            heading_levels[level] = []
                        heading_levels[level].append({
                            'text': para.text.strip(),
                            'index': i
                        })
                
                # 分析字体信息
                for run in para.runs:
                    if run.font.name:
                        structure['fonts_used'].add(run.font.name)
                
                structure['paragraph_analysis'].append(para_info)
            
            structure['headings'] = heading_levels
            structure['fonts_used'] = list(structure['fonts_used'])
            
            # 提取和分析图片
            images_info = self._extract_images_info(file_path)
            structure['images_info'] = images_info
            structure['images_count'] = len(images_info)
            
            # 水印分析
            structure['watermark_analysis'] = self._analyze_watermarks(images_info)
            
            return structure
            
        except Exception as e:
            return {'error': f'DOCX结构分析失败: {str(e)}'}
    
    def _analyze_pdf_structure(self, file_path: str) -> Dict[str, Any]:
        """分析PDF文档结构"""
        try:
            if fitz is None:
                return {'error': 'PyMuPDF未安装'}
            
            doc = fitz.open(file_path)
            
            structure = {
                'total_pages': len(doc),
                'total_text_blocks': 0,
                'fonts_used': set(),
                'headings': {},
                'page_analysis': [],
                'images_count': 0,
                'tables_detected': 0,
                'images_info': [],
                'watermark_analysis': {}
            }
            
            heading_levels = {}
            
            for page_num in range(len(doc)):
                page = doc[page_num]
                
                page_info = {
                    'page_number': page_num + 1,
                    'text_blocks': [],
                    'images_count': len(page.get_images()),
                    'fonts_on_page': set()
                }
                
                # 分析文本块
                blocks = page.get_text("dict")["blocks"]
                for block in blocks:
                    if "lines" in block:
                        for line in block["lines"]:
                            for span in line["spans"]:
                                font_info = f"{span.get('font', 'Unknown')}-{span.get('size', 0):.1f}"
                                structure['fonts_used'].add(font_info)
                                page_info['fonts_on_page'].add(font_info)
                                
                                text = span.get('text', '').strip()
                                if text:
                                    # 简单的标题检测（基于字体大小）
                                    font_size = span.get('size', 0)
                                    if font_size > 14 and len(text) < 200:  # 可能是标题
                                        level = self._estimate_heading_level(font_size)
                                        if level not in heading_levels:
                                            heading_levels[level] = []
                                        heading_levels[level].append({
                                            'text': text,
                                            'page': page_num + 1,
                                            'font_size': font_size
                                        })
                
                page_info['fonts_on_page'] = list(page_info['fonts_on_page'])
                structure['page_analysis'].append(page_info)
            
            structure['fonts_used'] = list(structure['fonts_used'])
            structure['headings'] = heading_levels
            
            # 提取和分析图片
            doc.close()  # 先关闭再重新打开进行图片分析
            images_info = self._extract_images_info(file_path)
            structure['images_info'] = images_info
            structure['images_count'] = len(images_info)
            
            # 水印分析
            structure['watermark_analysis'] = self._analyze_watermarks(images_info)
            
            return structure
            
        except Exception as e:
            return {'error': f'PDF结构分析失败: {str(e)}'}
    
    def _estimate_heading_level(self, font_size: float) -> int:
        """基于字体大小估算标题层级"""
        if font_size >= 20:
            return 1
        elif font_size >= 16:
            return 2
        elif font_size >= 14:
            return 3
        else:
            return 4
    
    def _extract_images_info(self, file_path: str) -> List[Dict[str, Any]]:
        """提取文档中的图片信息，包括水印检测"""
        images_info = []
        
        try:
            if self.document_type == 'pdf' and fitz:
                doc = fitz.open(file_path)
                for page_num in range(min(10, len(doc))):  # 处理前10页的图片
                    page = doc[page_num]
                    for img_index, img in enumerate(page.get_images()):
                        try:
                            xref = img[0]
                            pix = fitz.Pixmap(doc, xref)
                            if pix.n - pix.alpha < 4:  # 确保是RGB图像
                                img_data = pix.tobytes("png")
                                img_base64 = base64.b64encode(img_data).decode()
                                
                                # 水印检测
                                watermark_result = self._detect_watermark_from_pixmap(pix)
                                
                                img_info = {
                                    'page': page_num + 1,
                                    'index': img_index,
                                    'base64': img_base64,
                                    'width': pix.width,
                                    'height': pix.height,
                                    'size_kb': len(img_data) / 1024,
                                    'format': 'PNG',
                                    'watermark_detected': watermark_result['has_watermark'],
                                    'watermark_confidence': watermark_result['confidence'],
                                    'watermark_type': watermark_result['type']
                                }
                                images_info.append(img_info)
                            pix = None
                        except Exception as e:
                            # 记录图片处理错误，但继续处理其他图片
                            print(f"处理第{page_num + 1}页第{img_index}张图片时出错: {str(e)}")
                            continue
                doc.close()
                
            elif self.document_type == 'docx' and Document:
                # DOCX图片处理
                doc = Document(file_path)
                
                # 从关系中提取图片
                img_index = 0
                for rel in doc.part.rels.values():
                    if "image" in rel.target_ref:
                        try:
                            img_data = rel.target_part.blob
                            
                            # 转换为PIL图像进行分析
                            img_pil = Image.open(io.BytesIO(img_data))
                            img_base64 = base64.b64encode(img_data).decode()
                            
                            # 水印检测
                            watermark_result = self._detect_watermark_from_pil(img_pil)
                            
                            img_info = {
                                'index': img_index,
                                'base64': img_base64,
                                'width': img_pil.width,
                                'height': img_pil.height,
                                'size_kb': len(img_data) / 1024,
                                'format': img_pil.format or 'Unknown',
                                'watermark_detected': watermark_result['has_watermark'],
                                'watermark_confidence': watermark_result['confidence'],
                                'watermark_type': watermark_result['type']
                            }
                            images_info.append(img_info)
                            img_index += 1
                            
                        except Exception as e:
                            print(f"处理DOCX图片{img_index}时出错: {str(e)}")
                            continue
                        
        except Exception as e:
            print(f"图片提取失败: {str(e)}")
        
        return images_info[:20]  # 最多返回20张图片
    
    def _detect_watermark_from_pixmap(self, pixmap) -> Dict[str, Any]:
        """从PyMuPDF Pixmap检测水印"""
        try:
            # 转换为PIL图像进行分析
            img_data = pixmap.tobytes("png")
            img_pil = Image.open(io.BytesIO(img_data))
            return self._detect_watermark_from_pil(img_pil)
        except:
            return {'has_watermark': False, 'confidence': 0.0, 'type': 'unknown'}
    
    def _detect_watermark_from_pil(self, img_pil: Image.Image) -> Dict[str, Any]:
        """从PIL图像检测水印"""
        try:
            import numpy as np
            
            # 转换为RGB模式
            if img_pil.mode != 'RGB':
                img_pil = img_pil.convert('RGB')
            
            # 转换为numpy数组
            img_array = np.array(img_pil)
            
            # 水印检测策略
            result = {
                'has_watermark': False,
                'confidence': 0.0,
                'type': 'none'
            }
            
            # 1. 检查透明度（半透明水印）
            if img_pil.mode in ['RGBA', 'LA']:
                alpha_channel = img_array[:, :, -1] if img_pil.mode == 'RGBA' else img_array[:, :, 1]
                alpha_variation = np.std(alpha_channel)
                if alpha_variation > 30:  # 透明度变化较大，可能有水印
                    result['confidence'] += 0.3
                    result['type'] = 'transparent'
            
            # 2. 检查重复模式（重复水印）
            gray = np.mean(img_array, axis=2) if len(img_array.shape) == 3 else img_array
            
            # 简单的模式检测：检查图像的方差
            variance = np.var(gray)
            if variance < 500:  # 方差较小可能表示有重复模式
                result['confidence'] += 0.2
                if result['type'] == 'none':
                    result['type'] = 'pattern'
            
            # 3. 检查边缘区域（边框水印）
            height, width = gray.shape
            border_width = min(height, width) // 20
            
            if border_width > 5:
                top_border = gray[:border_width, :]
                bottom_border = gray[-border_width:, :]
                left_border = gray[:, :border_width]
                right_border = gray[:, -border_width:]
                
                # 检查边缘是否有特殊模式
                borders = [top_border, bottom_border, left_border, right_border]
                for border in borders:
                    if np.std(border) > 20:  # 边缘有变化
                        result['confidence'] += 0.1
                        if result['type'] == 'none':
                            result['type'] = 'border'
            
            # 4. 检查中心区域（中心水印）
            center_h, center_w = height // 2, width // 2
            quarter_h, quarter_w = height // 4, width // 4
            
            center_region = gray[center_h-quarter_h:center_h+quarter_h, 
                               center_w-quarter_w:center_w+quarter_w]
            
            if center_region.size > 0:
                center_std = np.std(center_region)
                overall_std = np.std(gray)
                
                # 如果中心区域的变化与整体不同，可能有水印
                if abs(center_std - overall_std) > 15:
                    result['confidence'] += 0.2
                    if result['type'] == 'none':
                        result['type'] = 'center'
            
            # 5. 检查颜色分布（彩色水印）
            if len(img_array.shape) == 3:
                # 计算各颜色通道的分布
                r_std = np.std(img_array[:, :, 0])
                g_std = np.std(img_array[:, :, 1])
                b_std = np.std(img_array[:, :, 2])
                
                # 如果某个通道特别突出，可能有彩色水印
                max_std = max(r_std, g_std, b_std)
                min_std = min(r_std, g_std, b_std)
                
                if max_std / (min_std + 1) > 2:  # 颜色分布不均
                    result['confidence'] += 0.15
                    if result['type'] == 'none':
                        result['type'] = 'colored'
            
            # 最终判断
            if result['confidence'] > 0.3:
                result['has_watermark'] = True
            
            # 限制置信度在0-1之间
            result['confidence'] = min(1.0, result['confidence'])
            
            return result
            
        except Exception as e:
            print(f"水印检测失败: {str(e)}")
            return {'has_watermark': False, 'confidence': 0.0, 'type': 'error'}
    
    def _analyze_watermarks(self, images_info: List[Dict[str, Any]]) -> Dict[str, Any]:
        """分析整个文档的水印情况"""
        if not images_info:
            return {'has_watermark': False, 'total_images': 0}
        
        watermark_count = sum(1 for img in images_info if img.get('watermark_detected', False))
        total_images = len(images_info)
        
        # 收集水印类型
        watermark_types = []
        total_confidence = 0
        
        for img in images_info:
            if img.get('watermark_detected', False):
                watermark_types.append(img.get('watermark_type', 'unknown'))
                total_confidence += img.get('watermark_confidence', 0)
        
        # 判断整个文档是否有水印
        has_watermark = watermark_count > 0
        avg_confidence = total_confidence / max(watermark_count, 1)
        
        # 确定主要水印类型
        if watermark_types:
            main_type = max(set(watermark_types), key=watermark_types.count)
        else:
            main_type = 'none'
        
        return {
            'has_watermark': has_watermark,
            'total_images': total_images,
            'watermark_images': watermark_count,
            'watermark_ratio': watermark_count / total_images if total_images > 0 else 0,
            'confidence': avg_confidence,
            'watermark_type': main_type,
            'watermark_types': list(set(watermark_types))
        }
    
    def _prepare_ai_analysis_data(self, analysis: Dict[str, Any]) -> Dict[str, Any]:
        """为AI分析准备结构化数据"""
        ai_data = {
            'document_summary': {},
            'structure_insights': {},
            'content_themes': [],
            'analysis_suggestions': []
        }
        
        # 文档摘要
        file_info = analysis['file_info']
        structure = analysis['structure_analysis']
        
        ai_data['document_summary'] = {
            'name': file_info['name'],
            'type': file_info['type'],
            'size_mb': file_info['size_mb'],
            'estimated_pages': structure.get('total_pages', analysis['preview_data'].get('page_count', 0)),
            'word_count': analysis['preview_data'].get('word_count', 0)
        }
        
        # 结构洞察
        if 'headings' in structure:
            heading_summary = {}
            for level, headings in structure['headings'].items():
                heading_summary[f'level_{level}'] = {
                    'count': len(headings),
                    'examples': [h['text'][:50] for h in headings[:3]]
                }
            ai_data['structure_insights']['headings'] = heading_summary
        
        if 'fonts_used' in structure:
            ai_data['structure_insights']['fonts_count'] = len(structure['fonts_used'])
            ai_data['structure_insights']['main_fonts'] = structure['fonts_used'][:5]
        
        # 分析建议
        ai_data['analysis_suggestions'] = [
            "文档用途分析：基于标题结构和内容判断文档类型（合同、报告、手册等）",
            "关键信息提取：识别重要的数据点、日期、金额、人名等",
            "结构化内容提取：按章节组织内容，便于后续分析",
            "关键词搜索：支持精确搜索和上下文提取"
        ]
        
        return ai_data
    
    def _prepare_search_capabilities(self, file_path: str) -> Dict[str, Any]:
        """准备搜索功能"""
        search_data = {
            'indexed_content': {},
            'search_ready': False
        }
        
        try:
            # 提取全文内容用于搜索
            full_text = ""
            
            if self.document_type == 'docx' and Document:
                doc = Document(file_path)
                full_text = "\n".join([para.text for para in doc.paragraphs])
            elif self.document_type == 'pdf' and fitz:
                doc = fitz.open(file_path)
                full_text = ""
                for page in doc:
                    full_text += page.get_text() + "\n"
                doc.close()
            
            if full_text:
                search_data['indexed_content'] = {
                    'full_text': full_text,
                    'paragraph_count': len(full_text.split('\n')),
                    'char_count': len(full_text)
                }
                search_data['search_ready'] = True
            
        except Exception as e:
            search_data['error'] = f"搜索索引准备失败: {str(e)}"
        
        return search_data
    
    def get_page_count(self, file_path: str = None) -> int:
        """
        获取文档页数
        
        Args:
            file_path: 文档路径，如果不提供则使用当前分析的文档
            
        Returns:
            文档页数
        """
        if file_path is None:
            file_path = self.document_path
        
        if not file_path or not os.path.exists(file_path):
            return 0
            
        doc_type = self._detect_document_type(file_path)
        
        if doc_type == 'pdf':
            try:
                if PYMUPDF_AVAILABLE:
                    doc = fitz.open(file_path)
                    page_count = len(doc)
                    doc.close()
                    return page_count
                elif PYPDF2_AVAILABLE:
                    with open(file_path, 'rb') as file:
                        pdf_reader = PyPDF2.PdfReader(file)
                        return len(pdf_reader.pages)
            except Exception as e:
                print(f"⚠️ 获取PDF页数失败: {e}")
                return 0
        
        elif doc_type == 'docx':
            try:
                if DOCX_AVAILABLE:
                    # DOCX没有明确的"页数"概念，这里估算
                    doc = Document(file_path)
                    paragraph_count = len(doc.paragraphs)
                    # 粗略估算：每页约15-20个段落
                    estimated_pages = max(1, paragraph_count // 18)
                    return estimated_pages
            except Exception as e:
                print(f"⚠️ 获取DOCX页数失败: {e}")
                return 0
        
        return 0
    
    def analyze_structure(self, file_path: str = None) -> Dict[str, Any]:
        """
        分析文档结构（为了兼容AI生成的代码）
        
        Args:
            file_path: 文档路径，如果不提供则使用当前分析的文档
            
        Returns:
            文档结构分析结果
        """
        if file_path is None:
            file_path = self.document_path
            
        if not file_path or not os.path.exists(file_path):
            return {}
            
        try:
            return self._analyze_document_structure(file_path)
        except Exception as e:
            print(f"⚠️ 结构分析失败: {e}")
            return {}
    
    def get_document_info(self, file_path: str = None) -> Dict[str, Any]:
        """
        获取文档基本信息
        
        Args:
            file_path: 文档路径，如果不提供则使用当前分析的文档
            
        Returns:
            文档基本信息
        """
        if file_path is None:
            file_path = self.document_path
            
        if not file_path or not os.path.exists(file_path):
            return {}
        
        return {
            'name': Path(file_path).name,
            'type': self._detect_document_type(file_path),
            'size_mb': round(os.path.getsize(file_path) / (1024 * 1024), 2),
            'page_count': self.get_page_count(file_path)
        }

    def search_keyword_context(self, keyword: str, context_lines: int = 3) -> List[Dict[str, Any]]:
        """
        搜索关键词并返回上下文
        
        Args:
            keyword: 要搜索的关键词
            context_lines: 上下文行数
            
        Returns:
            包含搜索结果和上下文的列表
        """
        if not self.analysis_result or 'search_capabilities' not in self.analysis_result:
            return []
        
        search_data = self.analysis_result['search_capabilities']
        if not search_data.get('search_ready'):
            return []
        
        full_text = search_data['indexed_content']['full_text']
        lines = full_text.split('\n')
        
        results = []
        for i, line in enumerate(lines):
            if keyword.lower() in line.lower():
                # 提取上下文
                start_idx = max(0, i - context_lines)
                end_idx = min(len(lines), i + context_lines + 1)
                
                context = '\n'.join(lines[start_idx:end_idx])
                
                results.append({
                    'line_number': i + 1,
                    'matched_line': line.strip(),
                    'context': context,
                    'keyword_position': line.lower().find(keyword.lower())
                })
        
        return results
    
    def generate_analysis_code(self, task_description: str) -> str:
        """
        根据任务描述生成智能分析代码
        
        Args:
            task_description: 任务描述
            
        Returns:
            生成的Python代码
        """
        if not self.analysis_result:
            return "# 请先分析文档"
        
        # 基础代码模板
        code_template = f'''
# 文档智能分析代码
# 文档: {self.analysis_result["file_info"]["name"]}
# 任务: {task_description}

from document_analyzer import DocumentAnalyzer
import json

# 初始化分析器
analyzer = DocumentAnalyzer()

# 分析文档
file_path = "{self.document_path}"
analysis_result = analyzer.analyze_document(file_path)

# 打印基本信息
print("=== 文档基本信息 ===")
file_info = analysis_result["file_info"]
print(f"文件名: {{file_info['name']}}")
print(f"类型: {{file_info['type']}}")
print(f"大小: {{file_info['size_mb']}} MB")

# 结构分析
print("\\n=== 文档结构分析 ===")
structure = analysis_result["structure_analysis"]
'''
        
        # 根据任务类型添加特定代码
        if "关键词" in task_description or "搜索" in task_description:
            # 提取可能的关键词
            potential_keywords = self._extract_potential_keywords(task_description)
            
            code_template += f'''
# 关键词搜索分析
keywords = {potential_keywords}

print("\\n=== 关键词搜索结果 ===")
for keyword in keywords:
    print(f"\\n搜索关键词: {{keyword}}")
    results = analyzer.search_keyword_context(keyword, context_lines=2)
    if results:
        for i, result in enumerate(results[:3], 1):  # 显示前3个结果
            print(f"  结果 {{i}}:")
            print(f"    行号: {{result['line_number']}}")
            print(f"    匹配行: {{result['matched_line'][:100]}}")
            print(f"    上下文:")
            print("    " + "\\n    ".join(result['context'].split('\\n')))
            print("-" * 50)
    else:
        print(f"  未找到关键词: {{keyword}}")
'''
        
        if "标题" in task_description or "结构" in task_description:
            code_template += '''
# 标题结构分析
if "headings" in structure:
    print("\\n=== 标题层级结构 ===")
    headings = structure["headings"]
    for level in sorted(headings.keys()):
        print(f"\\n{level}级标题 (共{len(headings[level])}个):")
        for heading in headings[level][:5]:  # 显示前5个
            print(f"  - {heading.get('text', heading)}")
'''
        
        if "字体" in task_description or "格式" in task_description:
            code_template += '''
# 字体使用分析
if "fonts_used" in structure:
    print("\\n=== 字体使用分析 ===")
    fonts = structure["fonts_used"]
    print(f"共使用了 {len(fonts)} 种字体:")
    for font in fonts[:10]:  # 显示前10种字体
        print(f"  - {font}")
'''
        
        code_template += '''
# 保存分析结果
output_file = f"{file_path}_analysis_result.json"
with open(output_file, 'w', encoding='utf-8') as f:
    json.dump(analysis_result, f, ensure_ascii=False, indent=2)

print(f"\\n✅ 分析完成！结果已保存到: {output_file}")
'''
        
        return code_template
    
    def _extract_potential_keywords(self, task_description: str) -> List[str]:
        """从任务描述中提取可能的关键词"""
        # 预定义的常见关键词
        common_keywords = [
            "合同编号", "合同号", "协议编号", "文档编号",
            "甲方", "乙方", "委托方", "受托方",
            "金额", "费用", "价格", "总价", "单价",
            "日期", "时间", "期限", "有效期",
            "责任", "义务", "权利", "条款",
            "签字", "盖章", "确认", "批准"
        ]
        
        # 从任务描述中提取引号内的内容
        quoted_keywords = re.findall(r'"([^"]*)"', task_description)
        quoted_keywords.extend(re.findall(r"'([^']*)'", task_description))
        quoted_keywords.extend(re.findall(r"「([^」]*)」", task_description))
        
        # 合并关键词
        all_keywords = quoted_keywords + common_keywords[:5]  # 取前5个常见关键词
        
        return list(set(all_keywords))[:10]  # 去重并限制在10个以内

def main():
    """测试函数"""
    analyzer = DocumentAnalyzer()
    
    # 测试文档分析
    test_file = "test_document.docx"  # 替换为实际文件路径
    if os.path.exists(test_file):
        result = analyzer.analyze_document(test_file)
        print("文档分析完成：")
        print(json.dumps(result, ensure_ascii=False, indent=2))
    else:
        print("测试文件不存在")

if __name__ == "__main__":
    main() 